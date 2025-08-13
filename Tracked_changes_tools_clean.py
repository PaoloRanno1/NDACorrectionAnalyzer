from __future__ import annotations

import json
import os
import re
from dataclasses import asdict, dataclass
from datetime import datetime, timezone
from typing import Any, Dict, Iterable, List, Optional, Tuple

from google import genai
try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    # Create dummy classes and functions for type hints
    class Document: pass
    class OxmlElement: pass
    class Table: pass
    class Paragraph: pass
    def qn(name): return name  # Dummy function
from copy import deepcopy
# =============================
# Data Models
# =============================


@dataclass
class CleanedFinding:
    """Represents a cleaned finding from the NDA review process."""
    id: int
    citation_clean: str
    suggested_replacement_clean: str


@dataclass
class RawFinding:
    """Represents a raw finding extracted from the reviewer's JSON output."""
    id: int  # 1-based index across all priorities (stable)
    priority: str  # "High Priority" | "Medium Priority" | "Low Priority"
    section: str
    issue: str
    problem: str
    citation: str
    suggested_replacement: str


# =============================
# Constants
# =============================

PROMPT_TEMPLATE = """
You are a precise contracts editor. Produce ONLY the JSON object specified in OUTPUT. No explanations, no code fences.

## Inputs:

-NDA TEXT (authoritative; do not alter):
{nda_text}
---

- RAW FINDING (verbatim JSON from previous step):
{raw_finding_json}
---

-ADDITIONAL GUIDANCE (may be empty):
{additional_info}
---

## Definitions:
- Exact substring = a contiguous sequence of characters that appears verbatim in NDA TEXT (same spelling, punctuation, capitalization, whitespace and dashes).
- Minimal sufficient span = the shortest exact substring that fully captures the problematic clause or phrase described in RAW FINDING.

## Tasks:
1) citation_clean:
Select one exact substring from NDA TEXT that:
   - Is contiguous, verbatim, and character-for-character identical to NDA TEXT.
   - Minimally captures the issue (don’t return entire sections if a sentence/phrase suffices).
   - Remove any truncations or ellipses from the input citation.
   - Removes any ellipses ... or truncations present in the RAW FINDING’s citation.
   - Preserves defined terms, numbering, symbols (including curly quotes/dashes as they appear in NDA TEXT), and capitalization.
Matching protocol (follow in order):
    - Try to locate the RAW FINDING’s citation exactly in NDA TEXT.
    - If not found, return empty string.
    - Never invent text. The chosen citation_clean must be verbatim from NDA TEXT.

2) suggested_replacement_clean:
Write one self-contained clause that can directly replace citation_clean
    - Align with issue and problem from RAW FINDING.
    - Use the RAW FINDING’s suggested_replacement as inspiration, but improve clarity/precision.
    - Look at the Examples of past replacements below to understand how to improve the suggested replacement.
    - Incorporate ADDITIONAL GUIDANCE if present (e.g., “less strict”, “retain carve-outs”, “Belgian law”).
    - Mirror the NDA’s style (formal drafting tone; keep defined terms consistent).
    - Concise and contract-ready. No commentary, notes, brackets like [ ], or rationale.
    - Stay close to the structure of citation_clean where reasonable (helps tracked-changes). Very important to try to preserve the original structure.
    - Avoid adding obligations beyond the policy intent (no scope creep).


## Output:
Return ONLY a JSON object in this form:
{{
  "id": <int>,
  "citation_clean": "<string>",
  "suggested_replacement_clean": "<string>"
}}

## Constraints:
- No explanations, notes, or commentary.
- No extra keys or fields.
- No code fences or markdown.
- "citation_clean" must be an exact substring from NDA TEXT. if not, shrink until it is.
- If you cannot ensure an exact match, choose a shorter exact match instead.
- If the RAW FINDING’s citation spans multiple non-contiguous parts in NDA TEXT, choose the single minimal contiguous part that carries the issue.

## Examples of past replacements
Follow these examples to improve the suggested replacement

- Issue: Policy 2 – Permitted Recipients
    - citation_clean: “…officers, directors, employees, professional advisers, providers of finance and auditors…”
    - suggested_replacement_clean: add shareholders and potential syndicate members while preserving original structure.

- Issue: Policy 3 – Return/Destruction
    - citation_clean: "on first written request by the Company, the Undersigned will within five business days at its option (i) return or (ii) destroy all Confidential Information and as far a reasonable practicable, permanently destroys all other data carriers, electronically, in writing or otherwise, carrying Confidential Information or making reference thereto"
    - suggested_replacement_clean: "on first written request by the Company, the Undersigned will within five business days at its option (i) return or (ii) destroy all Confidential Information and as far a reasonable practicable, permanently destroys all other data carriers, electronically, in writing or otherwise, carrying Confidential Information or making reference thereto 1 , except for (i) Confidential Information and elaborations thereof as have been saved to electronic carriers under automatic data archiving or data security procedures, (ii) reports, notes or other materials prepared by or on behalf of the Undersigned which incorporate Confidential Information (the “Secondary Information”), provided that the Secondary Information is kept confidential in accordance with the terms of this Undertaking, and (iii) Confidential Information which is required to be retained for the purposes of complying with any mandatory judicial, governmental, supervisory or regulatory order or for audit or compliance purposes;"

- Issue: POLICY 1: Handling of Unacceptable Clauses
    - citation_clean: : "If one of the parties fails to comply with the provisions of this agreement it is in default by operation of law and it forfeits an immediately due and payable penalty of EURO 10,000 per incident notwithstanding the right of the other party to claim all actual costs and damages and to invoke any other rights and remedies it may have pursuant to this agreement and applicable law."
    - suggested_replacement_clean: "If one of the parties fails to comply with the provisions of this agreement it is in default by operation of law and it forfeits an immediately due and payable penalty of the amount of which will be established by the courts notwithstanding the right of the other party to claim all actual costs and damages and to invoke any other rights and remedies it may have pursuant to this agreement and applicable law."

- Issue: POLICY 10: Governing Law & Jurisdiction
    - citation_clean: “This letter and any non-contractual obligations arising out of or in connection with this letter shall be governed by, and construed in accordance with, English law, and each party irrevocably submits to the exclusive jurisdiction of the English courts.”
    - suggested_replacement_clean: "This letter and any non-contractual obligations arising out of or in connection with this letter shall be governed by, and construed in accordance with, Belgian law, and each party irrevocably submits to the exclusive jurisdiction of the courts of Antwerp, division Antwerp."

- Issue: POLICY 11: Confidentiality Term
    - citation_clean: “the provisions of this Confidentiality Agreement shall terminate three years from that date of this Confidentiality Agreement, unless we receive prior written consent from Deloitte FA;”
    - suggested_replacement_clean: "the provisions of this Confidentiality Agreement shall terminate two years from that date of this Confidentiality Agreement, unless we receive prior written consent from Deloitte FA;"

"""


# =============================
# LLM Integration Functions
# =============================


def _call_gemini_json_prompt(prompt: str, model: str = "gemini-2.5-flash") -> Dict[str, Any]:
    """
    Calls the Gemini model with a plain text prompt, expecting a raw JSON object in the response.
    Handles basic extraction if extra whitespace is present.
    """
    client = genai.Client(api_key=os.environ["GOOGLE_API_KEY"])
    resp = client.models.generate_content(model=model, contents=prompt)
    text = resp.text.strip()

    # Extract the outermost JSON object, forgiving leading/trailing noise.
    # Use a simpler approach to find JSON-like content
    start_idx = text.find('{')
    if start_idx == -1:
        json_str = text
    else:
        # Find the matching closing brace
        brace_count = 0
        end_idx = start_idx
        for i, char in enumerate(text[start_idx:], start_idx):
            if char == '{':
                brace_count += 1
            elif char == '}':
                brace_count -= 1
                if brace_count == 0:
                    end_idx = i
                    break
        json_str = text[start_idx:end_idx+1]
    
    return json.loads(json_str)


def clean_findings_with_llm(
    nda_text: str,
    findings: List[RawFinding],
    additional_info_by_id: Optional[Dict[int, str]] = None,
    model: str = "gemini-2.5-pro",
) -> List[CleanedFinding]:
    """
    Cleans each RawFinding using an LLM to extract a verbatim citation substring and refine the suggested replacement.
    Validates that the cleaned citation is an exact substring of the NDA text.
    """
    additional_info_by_id = additional_info_by_id or {}
    results: List[CleanedFinding] = []

    for f in findings:
        raw_json = json.dumps(asdict(f), ensure_ascii=False)
        guidance = additional_info_by_id.get(f.id, "").strip()

        prompt = PROMPT_TEMPLATE.format(
            nda_text=nda_text,
            raw_finding_json=raw_json,
            additional_info=guidance,
        )

        try:
            obj = _call_gemini_json_prompt(prompt, model=model)
        except Exception as e:
            raise RuntimeError(f"LLM call failed for finding id={f.id}: {e}")

        # Basic schema validation
        for key in ("id", "citation_clean", "suggested_replacement_clean"):
            if key not in obj:
                raise ValueError(f"LLM output missing '{key}' for id={f.id}: {obj}")

        # Type normalization
        cid = int(obj["id"])
        citation_clean = str(obj["citation_clean"])
        sugg_clean = str(obj["suggested_replacement_clean"]).strip()

        # Validation: citation_clean MUST be a direct substring of nda_text
        if citation_clean not in nda_text:
            # Fallback: normalize whitespace and check again
            def _norm(s: str) -> str:
                return re.sub(r"\s+", " ", s).strip()

            if _norm(citation_clean) and _norm(citation_clean) in _norm(nda_text):
                pass  # Accept normalized match but retain original citation_clean
            else:
                raise ValueError(
                    f"[id={cid}] citation_clean is not an exact substring of NDA text.\n"
                    f"citation_clean: {citation_clean[:200]}...\n"
                    "Tip: Re-run with stronger guidance or shorten the expected span."
                )

        results.append(
            CleanedFinding(
                id=cid,
                citation_clean=citation_clean,
                suggested_replacement_clean=sugg_clean,
            )
        )

    return results


# =============================
# Document Processing Utilities
# =============================


def extract_text(docx_path: str) -> str:
    """Extracts all paragraph text from a DOCX file, joined by newlines."""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx library is not available. Cannot process DOCX files.")
    doc = Document(docx_path)
    return "\n".join(p.text for p in doc.paragraphs)


def flatten_findings(reviewer_json: Dict[str, Any]) -> List[RawFinding]:
    """
    Flattens findings from a reviewer's JSON structure into a list of RawFinding objects,
    assigning sequential IDs across priorities.
    """
    out: List[RawFinding] = []
    i = 1
    for prio in ("High Priority", "Medium Priority", "Low Priority"):
        for item in reviewer_json.get(prio, []):
            out.append(
                RawFinding(
                    id=i,
                    priority=prio,
                    section=str(item.get("section", "")),
                    issue=item.get("issue", ""),
                    problem=item.get("problem", ""),
                    citation=item.get("citation", "") or "",
                    suggested_replacement=item.get("suggested_replacement", "") or "",
                )
            )
            i += 1
    return out


def select_findings(
    findings: List[RawFinding],
    indices_to_keep: Optional[Iterable[int]] = None,
    where: Optional[callable] = None,
) -> List[RawFinding]:
    """
    Filters the list of findings based on indices or a predicate function.
    If neither is provided, returns all findings.
    """
    if indices_to_keep is not None:
        idxset = set(indices_to_keep)
        return [f for i, f in enumerate(findings) if i in idxset]
    if where is not None:
        return [f for f in findings if where(f)]
    return findings


def apply_edit_spec(items: List[RawFinding], edit_spec: Dict[str, Any]) -> List[RawFinding]:
    """
    Applies an edit specification to select and modify findings:
    - Accepts or discards based on sets.
    - Applies overrides to suggested replacements.
    - Supports citation hints for further processing.
    """
    accept_all_by_default = edit_spec.get("accept_all_by_default", False)
    accept_set = set(edit_spec.get("accept", []))
    discard_set = set(edit_spec.get("discard", []))
    overrides = edit_spec.get("overrides", {}) or {}

    selected: List[RawFinding] = []
    for it in items:
        keep = accept_all_by_default
        if it.id in accept_set:
            keep = True
        if it.id in discard_set:
            keep = False
        if it.id in overrides:  # Override implies keep
            keep = True
        if keep:
            # Apply overrides
            ov = overrides.get(it.id, {})
            sr = ov.get("suggested_replacement")
            if sr is not None:
                it = RawFinding(
                    id=it.id,
                    priority=it.priority,
                    section=it.section,
                    issue=it.issue,
                    problem=it.problem,
                    citation=it.citation,
                    suggested_replacement=sr,
                )
            # Optional citation hint (attached as private attribute)
            it._citation_hint = ov.get("citation_hint")  # type: ignore[attr-defined]
            selected.append(it)
    return selected


# =============================
# Tracked Changes Node Builders
# =============================


def _new_ins(author: str, dt_iso: str, text: str, change_id: str = "1"):
    """Creates an insertion (w:ins) XML element for tracked changes."""
    if not DOCX_AVAILABLE:
        return None
    ins = OxmlElement("w:ins")
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), dt_iso)
    ins.set(qn("w:id"), change_id)
    if text is not None:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        ins.append(r)
    return ins


def _new_del(author: str, dt_iso: str, text: str, change_id: str = "1"):
    """Creates a deletion (w:del) XML element for tracked changes."""
    if not DOCX_AVAILABLE:
        return None
    wdel = OxmlElement("w:del")
    wdel.set(qn("w:author"), author)
    wdel.set(qn("w:date"), dt_iso)
    wdel.set(qn("w:id"), change_id)
    r = OxmlElement("w:r")
    deltext = OxmlElement("w:delText")
    deltext.set(qn("xml:space"), "preserve")
    deltext.text = text
    r.append(deltext)
    wdel.append(r)
    return wdel


# =============================
# Character Map Utilities
# =============================

# Only define these constants if DOCX is available
if DOCX_AVAILABLE:
    TAB_TAG = qn("w:tab")
    BR_TAGS = {qn("w:br"), qn("w:cr")}
    T_TEXT = qn("w:t")
    RPR_TAG = qn("w:rPr")
else:
    TAB_TAG = "w:tab"
    BR_TAGS = {"w:br", "w:cr"}
    T_TEXT = "w:t"
    RPR_TAG = "w:rPr"

# Whitespace characters for boundary logic
_WS_CHARS = {" ", "\t", "\xa0", "\u2009", "\u200a", "\u200b", "\u202f"}  # space, tab, NBSP, thins, ZWSP


def _is_space_char(ch: str) -> bool:
    return ch in _WS_CHARS


def _display_char_for(child_tag: str, ch: Optional[str]) -> str:
    """Determines the visible character representation for matching purposes."""
    if child_tag == T_TEXT:
        if not ch:
            return ""
        return " " if _is_space_char(ch) else ch
    if child_tag == TAB_TAG or child_tag in BR_TAGS:
        return " "
    return ""


def _build_char_map(p: Paragraph) -> Tuple[List[Dict[str, Any]], str]:
    """
    Builds a character map for a paragraph, linking visible characters back to their XML elements.
    Returns the map and the logical text string.
    """
    char_map: List[Dict[str, Any]] = []
    for run in p.runs:
        r_el = run._r
        for child in r_el:
            tag = child.tag
            if tag == T_TEXT:
                txt = child.text or ""
                for i, ch in enumerate(txt):
                    disp = _display_char_for(tag, ch)
                    if disp != "":
                        char_map.append(
                            {"run": run, "child": child, "kind": "t", "idx": i, "raw": ch, "ch": disp}
                        )
            elif tag == TAB_TAG:
                char_map.append(
                    {"run": run, "child": child, "kind": "tab", "idx": None, "raw": "\t", "ch": " "}
                )
            elif tag in BR_TAGS:
                char_map.append(
                    {"run": run, "child": child, "kind": "br", "idx": None, "raw": "\n", "ch": " "}
                )
            # Ignore other inline nodes for matching
    logical = "".join(entry["ch"] for entry in char_map)
    return char_map, logical


def _paragraph_plain_text_logical(p: Paragraph) -> str:
    """Returns the logical plain text of a paragraph for matching."""
    return _build_char_map(p)[1]


def _find_all_matches(haystack: str, needle: str, ignore_case: bool = False) -> List[Tuple[int, int]]:
    """Finds all non-overlapping matches of needle in haystack."""
    if not needle:
        return []
    if ignore_case:
        haystack, needle = haystack.lower(), needle.lower()
    return [m.span() for m in re.finditer(re.escape(needle), haystack)]


def _remove_indices_from_textnode(child: OxmlElement, idxs: List[int]) -> None:
    """Removes specific indices from a w:t text node."""
    if child.tag != T_TEXT:
        return
    if not idxs:
        return
    txt = child.text or ""
    if not txt:
        return
    s = set(idxs)
    child.text = "".join(ch for j, ch in enumerate(txt) if j not in s)


def _is_run_visibly_empty(r_el: OxmlElement) -> bool:
    """Checks if a run has no visible content (text, tabs, breaks)."""
    has_visible = False
    for ch in r_el:
        tag = ch.tag
        if tag == RPR_TAG:
            continue
        if tag == T_TEXT:
            if (ch.text or "") != "":
                has_visible = True
                break
        elif tag == TAB_TAG or tag in BR_TAGS:
            has_visible = True
            break
        else:
            # Treat drawings/fields etc. as visible
            has_visible = True
            break
    return not has_visible


# =============================
# Boundary and Trim Helpers
# =============================

_PUNCT_RIGHT = set(",.;:!?)]}%»”’")
_PUNCT_LEFT = set("([{%«“‘")


def _is_space_entry(entry: Dict[str, Any]) -> bool:
    return entry["ch"] == " " and entry["kind"] in ("t", "tab", "br")


def _prev_char(char_map: List[Dict[str, Any]], idx: int) -> Optional[str]:
    i = idx - 1
    while i >= 0:
        ch = char_map[i]["ch"]
        if ch != "":
            return ch
        i -= 1
    return None


def _next_char(char_map: List[Dict[str, Any]], idx: int) -> Optional[str]:
    i = idx
    L = len(char_map)
    while i < L:
        ch = char_map[i]["ch"]
        if ch != "":
            return ch
        i += 1
    return None


def _trim_replacement_for_context(
    char_map: List[Dict[str, Any]], start: int, end: int, replacement: str
) -> str:
    """Trims the replacement text to avoid double spaces or misplaced punctuation based on context."""
    if not replacement:
        return replacement

    # Normalize internal whitespace to plain space
    rep = "".join(" " if _is_space_char(c) else c for c in replacement)

    prev_ch = _prev_char(char_map, start)
    next_ch = _next_char(char_map, end)

    # Avoid double spaces
    if prev_ch == " " and rep and rep[0] == " ":
        rep = rep[1:]
    if next_ch == " " and rep and rep[-1] == " ":
        rep = rep[:-1]

    # No space before right-side punctuation
    if next_ch in _PUNCT_RIGHT and rep.endswith(" "):
        rep = rep[:-1]

    # No space after left-side punctuation
    if prev_ch in _PUNCT_LEFT and rep.startswith(" "):
        rep = rep[1:]

    return rep


def _expand_bounds_for_whitespace(
    char_map: List[Dict[str, Any]], start: int, end: int
) -> Tuple[int, int]:
    """
    Expands the match bounds to include at most one adjacent whitespace character on each side.
    """
    L = len(char_map)
    if start > 0 and _is_space_entry(char_map[start - 1]):
        start -= 1
    if end < L and _is_space_entry(char_map[end]):
        end += 1
    return max(0, start), min(L, end)


# =============================
# Local Post-Clean Helpers
# =============================


def _remove_entry_space(entry: Dict[str, Any]) -> None:
    """Removes the space represented by a char_map entry."""
    if entry["kind"] == "t":
        _remove_indices_from_textnode(entry["child"], [entry["idx"]])
    elif entry["kind"] in ("tab", "br"):
        par = entry["child"].getparent()
        if par is not None:
            par.remove(entry["child"])


def _cleanup_paragraph_whitespace(p: Paragraph) -> None:
    """
    Cleans up stray whitespace in a paragraph after edits:
    - Removes trailing space at end of paragraph.
    - Removes space before right-side punctuation.
    - Collapses one instance of double-space.
    """
    char_map, _ = _build_char_map(p)
    if not char_map:
        return

    changed = False

    # 1) Trailing space at end of paragraph
    if char_map[-1]["ch"] == " ":
        _remove_entry_space(char_map[-1])
        changed = True

    # Rebuild if changed
    if changed:
        char_map, _ = _build_char_map(p)
        changed = False

    # 2) Space before right-side punctuation (e.g., "word ,")
    for i in range(1, len(char_map)):
        if char_map[i]["ch"] in _PUNCT_RIGHT and char_map[i - 1]["ch"] == " ":
            _remove_entry_space(char_map[i - 1])
            changed = True
            break

    if changed:
        char_map, _ = _build_char_map(p)
        changed = False

    # 3) Collapse one double-space
    for i in range(len(char_map) - 1):
        if char_map[i]["ch"] == " " and char_map[i + 1]["ch"] == " ":
            _remove_entry_space(char_map[i])
            changed = True
            break

    # 4) Drop now-empty runs
    if changed:
        for run in list(p.runs):
            r_el = run._r
            if _is_run_visibly_empty(r_el):
                par = r_el.getparent()
                if par is not None:
                    par.remove(r_el)


# =============================
# Core Paragraph Operations
# =============================


def _apply_match_to_paragraph(
    p: Paragraph,
    start: int,
    end: int,
    replacement_text: str,
    author: str,
    dt_iso: str,
    change_id: str,
) -> bool:
    """
    Applies a tracked change (deletion and insertion) to a matched span in a paragraph.
    Handles boundary whitespace and cleanup.
    Returns True if the change was applied successfully.
    """
    if start >= end:
        return False

    char_map, _ = _build_char_map(p)
    if not char_map or end > len(char_map):
        return False

    # Expand to swallow boundary whitespace
    start, end = _expand_bounds_for_whitespace(char_map, start, end)
    if start >= end:
        return False

    # Trim replacement based on context
    replacement_text = _trim_replacement_for_context(char_map, start, end, replacement_text or "")

    # Build deleted text (logical characters)
    del_text = "".join(entry["ch"] for entry in char_map[start:end])

    # Insert change markers before editing
    first_anchor_run_el = char_map[start]["run"]._r
    parent = first_anchor_run_el.getparent()
    del_el = _new_del(author, dt_iso, del_text, change_id)
    ins_el = _new_ins(author, dt_iso, replacement_text, change_id) if replacement_text != "" else None

    if parent is not None:
        first_anchor_run_el.addprevious(del_el)
        if ins_el is not None:
            first_anchor_run_el.addprevious(ins_el)
    else:
        # Fallback
        p._p.append(del_el)
        if ins_el is not None:
            p._p.append(ins_el)

    # Delete selected characters/tabs/breaks
    remove_indices_by_child: Dict[Any, List[int]] = {}
    special_children_to_remove = []

    for i in range(start, end):
        entry = char_map[i]
        if entry["kind"] == "t":
            remove_indices_by_child.setdefault(entry["child"], []).append(entry["idx"])
        elif entry["kind"] in ("tab", "br"):
            if entry["child"] not in special_children_to_remove:
                special_children_to_remove.append(entry["child"])

    for child, idxs in remove_indices_by_child.items():
        _remove_indices_from_textnode(child, idxs)

    for ch_el in special_children_to_remove:
        par = ch_el.getparent()
        if par is not None:
            par.remove(ch_el)

    # Drop empty runs
    for run in list(p.runs):
        r_el = run._r
        if _is_run_visibly_empty(r_el):
            par = r_el.getparent()
            if par is not None:
                par.remove(r_el)

    # Local cleanup
    _cleanup_paragraph_whitespace(p)
    return True


# =============================
# Paragraph Iteration (Body + Tables)
# =============================


def _iter_paragraphs_in_document(doc: Document) -> Iterable[Paragraph]:
    """Yields all paragraphs in the document, including those in tables."""
    # Body paragraphs
    for p in doc.paragraphs:
        yield p
    # Table paragraphs
    for t in doc.tables:
        yield from _iter_paragraphs_in_table(t)


def _iter_paragraphs_in_table(table: Table) -> Iterable[Paragraph]:
    """Yields all paragraphs in a table, recursing into nested tables."""
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for t in cell.tables:
                yield from _iter_paragraphs_in_table(t)


# =============================
# Public API
# =============================


def apply_cleaned_findings_to_docx(
    input_docx: str,
    cleaned_findings: List[CleanedFinding],
    output_docx: str,
    author: str = "AI Reviewer",
    ignore_case: bool = False,
    skip_if_same: bool = True,
) -> int:
    """
    Applies tracked changes to a DOCX file based on cleaned findings.
    Matches are exact and contiguous within paragraphs (body and tables).
    Returns the number of changes applied.
    """
    doc = Document(input_docx)
    dt_iso = datetime.now(timezone.utc).isoformat(timespec="seconds")

    # Prepare worklist
    work: List[Tuple[int, str, str]] = []
    for f in cleaned_findings:
        citation = (f.citation_clean or "").strip()
        replacement = f.suggested_replacement_clean or ""
        if not citation:
            continue
        if skip_if_same and citation == replacement:
            continue
        work.append((f.id, citation, replacement))

    applied = 0
    change_counter = 1

    for _, citation, replacement in work:
        # Process every paragraph
        for p in _iter_paragraphs_in_document(doc):
            logical = _paragraph_plain_text_logical(p)
            if not logical:
                continue

            matches = _find_all_matches(logical, citation, ignore_case=ignore_case)
            if not matches:
                continue

            # Apply from end to preserve indices
            for start, end in reversed(matches):
                ok = _apply_match_to_paragraph(
                    p=p,
                    start=start,
                    end=end,
                    replacement_text=replacement,
                    author=author,
                    dt_iso=dt_iso,
                    change_id=str(change_counter),
                )
                if ok:
                    applied += 1
                    change_counter += 1

    # Final cleanup sweep
    for p in _iter_paragraphs_in_document(doc):
        _cleanup_paragraph_whitespace(p)

    doc.save(output_docx)
    return applied


####
####
'''
Function to output new docx file
'''
def _get(obj: Any, key: str, default: str = "") -> str:
    """Get attribute or dict item from CleanedFinding-like objects."""
    if hasattr(obj, key):
        return getattr(obj, key) or default
    if isinstance(obj, dict):
        return obj.get(key, default) or default
    return default

# =============================
# Character / XML Utilities
# =============================

# Constants are already defined above

def _is_space_char(ch: str) -> bool:
    return ch in _WS_CHARS

def _display_char_for(child_tag: str, ch: Optional[str]) -> str:
    """Visible character for matching purposes."""
    if child_tag == T_TEXT:
        if not ch:
            return ""
        return " " if _is_space_char(ch) else ch
    if child_tag == TAB_TAG or child_tag in BR_TAGS:
        return " "
    return ""

def _build_char_map(p: Paragraph) -> Tuple[List[Dict[str, Any]], str]:
    """
    Build a map of visible characters in a paragraph back to (run, child, idx) entries.
    Returns (map, logical_text).
    """
    char_map: List[Dict[str, Any]] = []
    for run in p.runs:
        r_el = run._r  # XML <w:r>
        for child in r_el:
            tag = child.tag
            if tag == T_TEXT:
                txt = child.text or ""
                for i, ch in enumerate(txt):
                    disp = _display_char_for(tag, ch)
                    if disp != "":
                        char_map.append({"run": run, "child": child, "kind": "t", "idx": i, "raw": ch, "ch": disp})
            elif tag == TAB_TAG:
                char_map.append({"run": run, "child": child, "kind": "tab", "idx": None, "raw": "\t", "ch": " "})
            elif tag in BR_TAGS:
                char_map.append({"run": run, "child": child, "kind": "br", "idx": None, "raw": "\n", "ch": " "})
            else:
                # Treat other inline nodes as invisible for matching purposes
                pass
    logical = "".join(entry["ch"] for entry in char_map)
    return char_map, logical

def _paragraph_plain_text_logical(p: Paragraph) -> str:
    return _build_char_map(p)[1]

def _find_all_matches(haystack: str, needle: str, ignore_case: bool = False) -> List[Tuple[int, int]]:
    """All non-overlapping exact matches."""
    if not needle:
        return []
    if ignore_case:
        haystack, needle = haystack.lower(), needle.lower()
    return [m.span() for m in re.finditer(re.escape(needle), haystack)]

def _remove_indices_from_textnode(child: OxmlElement, idxs: List[int]) -> None:
    """Remove characters at zero-based indices from a <w:t> node."""
    if child.tag != T_TEXT or not idxs:
        return
    txt = child.text or ""
    if not txt:
        return
    s = set(idxs)
    child.text = "".join(ch for j, ch in enumerate(txt) if j not in s)

def _is_run_visibly_empty(r_el: OxmlElement) -> bool:
    """True if run has no visible content (text, tabs, breaks, drawings, etc.)."""
    has_visible = False
    for ch in r_el:
        tag = ch.tag
        if tag == RPR_TAG:
            continue
        if tag == T_TEXT:
            if (ch.text or "") != "":
                has_visible = True
                break
        elif tag == TAB_TAG or tag in BR_TAGS:
            has_visible = True
            break
        else:
            # other nodes (e.g., fields, drawings) -> visible
            has_visible = True
            break
    return not has_visible

# Punctuation context sets for trimming
_PUNCT_RIGHT = set(",.;:!?)]}%»”’")
_PUNCT_LEFT = set("([{%«“‘")

def _prev_char(char_map: List[Dict[str, Any]], idx: int) -> Optional[str]:
    i = idx - 1
    while i >= 0:
        ch = char_map[i]["ch"]
        if ch != "":
            return ch
        i -= 1
    return None

def _next_char(char_map: List[Dict[str, Any]], idx: int) -> Optional[str]:
    i = idx
    L = len(char_map)
    while i < L:
        ch = char_map[i]["ch"]
        if ch != "":
            return ch
        i += 1
    return None

def _is_space_entry(entry: Dict[str, Any]) -> bool:
    return entry["ch"] == " " and entry["kind"] in ("t", "tab", "br")

def _expand_bounds_for_whitespace(char_map: List[Dict[str, Any]], start: int, end: int) -> Tuple[int, int]:
    """Optionally swallow one adjacent whitespace on each side."""
    L = len(char_map)
    if start > 0 and _is_space_entry(char_map[start - 1]):
        start -= 1
    if end < L and _is_space_entry(char_map[end]):
        end += 1
    return max(0, start), min(L, end)

def _trim_replacement_for_context(char_map: List[Dict[str, Any]], start: int, end: int, replacement: str) -> str:
    """Avoid double-spaces / stray spaces around punctuation."""
    if not replacement:
        return replacement
    rep = "".join(" " if _is_space_char(c) else c for c in replacement)

    prev_ch = _prev_char(char_map, start)
    next_ch = _next_char(char_map, end)

    # Avoid double space at both ends
    if prev_ch == " " and rep and rep[0] == " ":
        rep = rep[1:]
    if next_ch == " " and rep and rep[-1] == " ":
        rep = rep[:-1]

    # Space rules around punctuation
    if next_ch in _PUNCT_RIGHT and rep.endswith(" "):
        rep = rep[:-1]
    if prev_ch in _PUNCT_LEFT and rep.startswith(" "):
        rep = rep[1:]

    return rep

def _cleanup_paragraph_whitespace(p: Paragraph) -> None:
    """Small cleanup pass: trailing space, space before punctuation, one double-space, and drop empty runs."""
    char_map, _ = _build_char_map(p)
    if not char_map:
        return

    changed = False

    # 1) Trailing space
    if char_map[-1]["ch"] == " ":
        _remove_entry_space(char_map[-1])
        changed = True

    if changed:
        char_map, _ = _build_char_map(p)
        changed = False

    # 2) "word ," -> remove space before punctuation
    for i in range(1, len(char_map)):
        if char_map[i]["ch"] in _PUNCT_RIGHT and char_map[i - 1]["ch"] == " ":
            _remove_entry_space(char_map[i - 1])
            changed = True
            break

    if changed:
        char_map, _ = _build_char_map(p)
        changed = False

    # 3) Collapse one double-space
    for i in range(len(char_map) - 1):
        if char_map[i]["ch"] == " " and char_map[i + 1]["ch"] == " ":
            _remove_entry_space(char_map[i])
            changed = True
            break

    # 4) Drop now-empty runs
    if changed:
        for run in list(p.runs):
            r_el = run._r
            if _is_run_visibly_empty(r_el):
                par = r_el.getparent()
                if par is not None:
                    par.remove(r_el)

def _remove_entry_space(entry: Dict[str, Any]) -> None:
    """Remove a space represented by a char_map entry."""
    if entry["kind"] == "t":
        _remove_indices_from_textnode(entry["child"], [entry["idx"]])
    elif entry["kind"] in ("tab", "br"):
        par = entry["child"].getparent()
        if par is not None:
            par.remove(entry["child"])

# =============================
# Paragraph Iteration (Body + Tables)
# =============================

def _iter_paragraphs_in_table(table: Table) -> Iterable[Paragraph]:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for t in cell.tables:  # nested
                yield from _iter_paragraphs_in_table(t)

def _iter_paragraphs_in_document(doc: Document) -> Iterable[Paragraph]:
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        yield from _iter_paragraphs_in_table(t)

# =============================
# Core: In-place (non-tracked) replacement
# =============================

def _insert_text_before_run(run_el: OxmlElement, text: str) -> None:
    """
    Insert a new <w:r><w:t>text</w:t></w:r> immediately BEFORE the given run element.
    Copies rPr from the anchor run if present to preserve formatting.
    """
    if text == "":
        return
    new_r = OxmlElement("w:r")
    # copy formatting
    if run_el.rPr is not None:
        new_r.append(deepcopy(run_el.rPr))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_r.append(t)
    run_el.addprevious(new_r)

def _apply_plain_replacement_to_paragraph(
    p: Paragraph,
    start: int,
    end: int,
    replacement_text: str,
) -> bool:
    """
    Replace [start:end] (logical char indices) in 'p' with 'replacement_text' WITHOUT tracked changes.
    Keeps the paragraph/run structure intact except for the target span and a new run inserted for the replacement.
    """
    if start >= end:
        return False

    char_map, _ = _build_char_map(p)
    if not char_map or end > len(char_map):
        return False

    # Swallow at most one adjacent whitespace on each side to avoid dangling spaces
    start, end = _expand_bounds_for_whitespace(char_map, start, end)
    if start >= end:
        return False

    # Context-aware trimming of replacement
    replacement_text = _trim_replacement_for_context(char_map, start, end, replacement_text or "")

    # If no-op replacement (empty and removing only spaces that are already trimmed to nothing) is fine
    first_anchor_run_el = char_map[start]["run"].__getattribute__("_r")  # <w:r>

    # 1) Delete selected characters/tabs/breaks
    remove_indices_by_child: Dict[Any, List[int]] = {}
    special_children_to_remove = []

    for i in range(start, end):
        entry = char_map[i]
        if entry["kind"] == "t":
            remove_indices_by_child.setdefault(entry["child"], []).append(entry["idx"])
        elif entry["kind"] in ("tab", "br"):
            if entry["child"] not in special_children_to_remove:
                special_children_to_remove.append(entry["child"])

    for child, idxs in remove_indices_by_child.items():
        _remove_indices_from_textnode(child, idxs)

    for ch_el in special_children_to_remove:
        par = ch_el.getparent()
        if par is not None:
            par.remove(ch_el)

    # 2) Insert the replacement text at the position of the first removed char
    if replacement_text:
        _insert_text_before_run(first_anchor_run_el, replacement_text)

    # 3) Drop empty runs and do a small whitespace cleanup
    for run in list(p.runs):
        r_el = run._r
        if _is_run_visibly_empty(r_el):
            par = r_el.getparent()
            if par is not None:
                par.remove(r_el)

    _cleanup_paragraph_whitespace(p)
    return True

# =============================
# Public API
# =============================

def replace_cleaned_findings_in_docx(
    input_docx: str,
    cleaned_findings: List[Any],
    output_docx: str,
    *,
    ignore_case: bool = False,
    skip_if_same: bool = True,
) -> int:
    """
    Replace each finding.citation_clean with finding.suggested_replacement_clean across the document
    (body + tables), WITHOUT tracked changes.

    The document's structure (styles, numbering, tables, headers/footers content that includes paragraphs, etc.)
    remains intact. Matches are exact, contiguous, and paragraph-local.

    Returns the number of replacements applied (each match counts as one).
    """
    doc = Document(input_docx)

    # Build worklist
    work: List[Tuple[int, str, str]] = []
    for f in cleaned_findings:
        citation = (_get(f, "citation_clean", "").strip())
        replacement = _get(f, "suggested_replacement_clean", "")
        if not citation:
            continue
        if skip_if_same and citation == replacement:
            continue
        # We don't actually need the id here, but keep the tuple shape similar to your previous API
        fid = int(_get(f, "id", "0") or 0)
        work.append((fid, citation, replacement))

    applied = 0

    # Iterate paragraphs and apply replacements (from the back to keep indices valid)
    for _, citation, replacement in work:
        for p in _iter_paragraphs_in_document(doc):
            logical = _paragraph_plain_text_logical(p)
            if not logical:
                continue

            matches = _find_all_matches(logical, citation, ignore_case=ignore_case)
            if not matches:
                continue

            for start, end in reversed(matches):
                ok = _apply_plain_replacement_to_paragraph(
                    p=p,
                    start=start,
                    end=end,
                    replacement_text=replacement,
                )
                if ok:
                    applied += 1

    # Final cleanliness pass
    for p in _iter_paragraphs_in_document(doc):
        _cleanup_paragraph_whitespace(p)

    doc.save(output_docx)
    return applied